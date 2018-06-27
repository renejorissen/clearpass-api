#!/usr/bin/env python3
#------------------------------------------------------------------------------
#
# Script to import GuestUsers from a CSV file
# The scripts creates a GuestUser per week of the year
#
#------------------------------------------------------------------------------

import requests
import json
from configparser import ConfigParser
import os
import urllib3
import csv
import pprint
import string
import random
import time
from docx import Document
from docx.shared import Inches
import datetime

urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

# configuration file parameters
params = os.path.join(os.path.dirname(__file__), "../config/params.cfg")
config = ConfigParser()
config.read(params)

clearpass_fqdn = config.get('ClearPass', 'clearpass_fqdn')
oauth_grant_type = config.get('OAuth2', 'grant_type')
oauth_client_id = config.get('OAuth2', 'client_id')
oauth_client_secret = config.get('OAuth2', 'client_secret')
oauth_username = config.get('OAuth2', 'username')
oauth_password = config.get('OAuth2', 'password')


# validate config
def check_config(clearpass_fqdn, oauth_grant_type, oauth_client_id, oauth_client_secret, oauth_username, oauth_password):
    """Validate the OAuth 2.0 configuration from the params.cfg file."""

    if not clearpass_fqdn:
        print('Error: ClearPass FQDN must be defined in config file (config/params.cfg)')
        exit(1)
    if not oauth_grant_type:
        print('Error: grant_type must be defined in config file (config/params.cfg)')
        exit(1)
    if not oauth_client_id:
        print('Error: client_id must be defined in config file (config/params.cfg)')
        exit(1)
    if oauth_grant_type == "password" and (not oauth_username or not oauth_password):
        print('Error: username and password must be defined in config file for password grant type (config/params.cfg)')
        exit(1)


####################################################################
###  AUTHENTICATION TO CLEARPASS
####################################################################

def get_access_token(clearpass_fqdn, oauth_grant_type, oauth_client_id, oauth_client_secret, oauth_username, oauth_password):
    """Get OAuth 2.0 access token with config from params.cfg"""

    url = "https://" + clearpass_fqdn + "/api/oauth"

    headers = {'Content-Type':'application/json'}

    # grant_type: password
    if oauth_grant_type == "password":
        payload = {'grant_type':oauth_grant_type, 'username':oauth_username, 'password':oauth_password, 'client_id':oauth_client_id, 'client_secret':oauth_client_secret}
        #print(payload)
        try:
            r = requests.post(url, headers=headers, json=payload, verify=False, timeout=2)
            r.raise_for_status()
        except Exception as e:
            print(e)
            exit(1)

        json_response = json.loads(r.text)

        return json_response

    # grant_type: password   public client
    if oauth_grant_type == "password" and not oauth_client_secret:
        payload = {'grant_type':oauth_grant_type, 'username':oauth_username, 'password':oauth_password, 'client_id':oauth_client_id}

        try:
            r = requests.post(url, headers=headers, json=payload, verify=False, timeout=2)
            r.raise_for_status()
        except Exception as e:
            print(e)
            exit(1)

        json_response = json.loads(r.text)

        return json_response

    # grant_type: client_credentials
    if oauth_grant_type == "client_credentials":
        payload = {'grant_type': oauth_grant_type, 'client_id': oauth_client_id, 'client_secret': oauth_client_secret}

        try:
            r = requests.post(url, headers=headers, json=payload, verify=False, timeout=2)
            r.raise_for_status()
        except Exception as e:
            print(e)
            exit(1)

        json_response = json.loads(r.text)

        return json_response

def get_api_role(clearpass_fqdn, token_type, access_token):
    """Get the current ClearPass operator profile name"""

    url = "https://" + clearpass_fqdn + "/api/oauth/me"

    headers = {'Content-Type':'application/json', "Authorization": "{} {}".format(token_type, access_token)}

    try:
        r = requests.get(url, headers=headers, verify=False, timeout=2)
        r.raise_for_status()
    except Exception as e:
        print(e)
        exit(1)

    json_response = json.loads(r.text)

    return json_response


def get_privs(clearpass_fqdn, token_type, access_token):
    """Get the current access privileges"""

    url = "https://" + clearpass_fqdn + "/api/oauth/privileges"

    headers = {'Content-Type':'application/json', "Authorization": "{} {}".format(token_type, access_token)}

    try:
        r = requests.get(url, headers=headers, verify=False, timeout=2)
        r.raise_for_status()
    except Exception as e:
        print(e)
        exit(1)

    json_response = json.loads(r.text)

    return json_response


check_config(clearpass_fqdn, oauth_grant_type, oauth_client_id, oauth_client_secret, oauth_username, oauth_password)

token_response = get_access_token(clearpass_fqdn, oauth_grant_type, oauth_client_id, oauth_client_secret, oauth_username, oauth_password)
access_token = token_response['access_token']
token_type = token_response['token_type']
token_expires_in = token_response['expires_in']
scope = token_response['scope']

get_api_role_response = get_api_role(clearpass_fqdn, token_type, access_token)
api_role = get_api_role_response['info']

get_privs_response = get_privs(clearpass_fqdn, token_type, access_token)
api_privs = get_privs_response['privileges']


####################################################################
###  ADD A USER
####################################################################

print("")
print("LET US ADD THE GUEST USER(S)")
print("============================")
print("")

#####################################################################
### IMPORT THE CSV FILE
### CSV FILE CREATED IN EXCEL 2016 AND SAVED AS "CSV (Comma delimited)(*.csv)

### WHEN OPENED WITH NOTEPAD THE CONTENT IS
###
###username;email;role_id;enabled;start_time;expire_time;simultaneous_use
###gast@week26;gast@week26;2;TRUE;1529877600;1530482400;0
###gast@week27;gast@week27;2;TRUE;1530482400;1531087200;0
###
### EPOCH value start_time is calculated via time_epoch.py
### EPOCH difference for a week = 604800
###
### PASSWORD IS GENERATED VIA PASSWORD GENERATOR
#######################################################################

with open('guests.csv', 'r') as file:
    reader = csv.DictReader(file, delimiter=';')
    user_list = []
    for line in reader:
        user_list.append(line)

        ####################################################################
        ###  PASSWORD GENERATOR FOR 15 CHARACTERS
        ####################################################################
        uppercase = (string.ascii_uppercase)
        lowercase = (string.ascii_lowercase)
        number = (string.digits)
        symbols = (string.punctuation)

        very_simple = lowercase + number
        simple = uppercase + lowercase + number
        strong = uppercase + lowercase + number + symbols

        generated_password = "".join(random.sample(very_simple, 6))

        ####################################################################
        ### DEFINE BASE URL FOR ADDING LOCAL USERS
        ####################################################################
        url = "https://" + clearpass_fqdn + "/api/guest"

        ####################################################################
        ### DEFINE THE HEADER PARAMETERS
        ####################################################################
        headers = {'Accept': 'application/json', "Authorization": "{} {}".format(token_type, access_token)}

        ####################################################################
        ### DEFINE THE PAYLOAD
        ####################################################################
        payload = {'email': line['email'], 'password': generated_password, 'username': line['username'], 'role_id': line['role_id'], 'simultaneous_use': line['simultaneous_use'],
                   'enabled': line['enabled'], 'start_time': line['start_time'], 'expire_time': line['expire_time']}
        #print(payload)

        ####################################################################
        ### ADDING THE USER(S)
        ####################################################################
        while True:
            post_user = requests.post(url, headers=headers, json=payload, verify=False, timeout=2)
            if post_user.status_code == 201:
                print("USER {} WAS CREATED SUCCESSFULLY".format(line['username']))
                print("INFO WILL BE WRITTEN TO WORD DOCUMENT IN GUEST_PASSES DIRECTORY: {}.docx".format(line['description']))
                print("")

                ####################################################################
                ### CONVERT EPOCH TO HUMAN READABLE FORMAT
                ####################################################################
                timestamp = int(line['start_time'])
                value = datetime.datetime.fromtimestamp(timestamp)
                activated_on = value.strftime('%Y-%m-%d %H:%M:%S')

                timestamp = int(line['expire_time'])
                value = datetime.datetime.fromtimestamp(timestamp)
                expired_on = value.strftime('%Y-%m-%d %H:%M:%S')

                ####################################################################
                ### PRINT OUTPUT TO WORD FILE
                ####################################################################
                document = Document()

                document.add_picture('guest_passes/logo_4ip.jpg', width=Inches(4.00))
                document.add_heading('Gast gebruiker', 0)

                p = document.add_paragraph('Onderstaand worden de account details voor deze week weergegeven.')

                table = document.add_table(rows=1, cols=4)
                table.style = 'TableGrid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Gebruiker'
                hdr_cells[1].text = 'Wachtwoord'
                hdr_cells[2].text = 'Geactiveerd op'
                hdr_cells[3].text = 'Verloopt op'
                cells = table.add_row().cells
                cells[0].text = line['username']
                cells[1].text = generated_password
                cells[2].text = activated_on
                cells[3].text = expired_on

                document.save('guest_passes/{} 2018.docx'.format(line['description']))
                break
            else:
                print("SOMETHING WENT WRONG! ERROR CODE = {}".format(post_user.status_code))
                print("")
                print("401 - Unauthorized")
                print("403 - Forbidden")
                print("406 - Not Acceptable")
                print("415 - Unsupported Media Type")
                print("422 - Unprocessable Entity")
                print("")
                print("LETS TRY AGAIN")
                print("")
                break
print("GOODBYE!!")