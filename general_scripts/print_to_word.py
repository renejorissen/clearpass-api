from docx import Document
from docx.shared import Inches

document = Document()

document.add_picture('../guests/guest_passes/logo_gsg.jpg', width=Inches(4.00))
document.add_heading('Gast gebruiker', 0)

p = document.add_paragraph('Onderstaand worden de account details voor deze week weergegeven.')
#p.add_run('bold').bold = True
#p.add_run(' and some ')
#p.add_run('italic.').italic = True

#document.add_heading('Heading, level 1', level=1)
#document.add_paragraph('Intense quote', style='IntenseQuote')

#document.add_paragraph(
#    'first item in unordered list', style='ListBullet'
#)
#document.add_paragraph(
#    'first item in ordered list', style='ListNumber'
#)



table = document.add_table(rows=1, cols=2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'


document.add_page_break()
username = "test_document"
document.save('../guests/guest_passes/{}.docx'.format(username))